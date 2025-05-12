import os
import docx
from .base_loader import BaseLoader

class DOCXLoader(BaseLoader):
    """
    Loads text content from DOCX files using python-docx library.
    """
    
    def load_document(self, file_path):
        """
        Extract text from a DOCX document.
        
        Args:
            file_path (str): Path to the DOCX file.
            
        Returns:
            str: The extracted text content from the DOCX.
        """
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return self.clean_text("\n".join(full_text))
        
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX '{file_path}': {str(e)}")
            
    def load_document_with_metadata(self, file_path):
        """
        Extract text and metadata from a DOCX document.
        
        Args:
            file_path (str): Path to the DOCX file.
            
        Returns:
            tuple: (text, metadata) where text is the extracted content and 
                   metadata is a dictionary with file information.
        """
        try:
            # Get file metadata
            file_stats = os.stat(file_path)
            file_name = os.path.basename(file_path)
            
            # Load the document
            doc = docx.Document(file_path)
            
            # Extract metadata
            metadata = {
                "file_name": file_name,
                "file_size": file_stats.st_size,
                "created_time": file_stats.st_ctime,
                "modified_time": file_stats.st_mtime,
                "file_extension": ".docx",
                "author": doc.core_properties.author if hasattr(doc.core_properties, 'author') else "",
                "title": doc.core_properties.title if hasattr(doc.core_properties, 'title') else "",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            # Get text content
            text = self.load_document(file_path)
            
            return text, metadata
        
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX '{file_path}': {str(e)}") 